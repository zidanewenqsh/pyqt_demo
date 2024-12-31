import sys
from PyQt5.QtWidgets import QApplication, QMainWindow, QFileDialog, QMessageBox
from PyQt5.QtGui import QStandardItemModel, QStandardItem
from Ui_wanlihong31 import Ui_MainWindow
from docx import Document


class DocumentGenerationTool(QMainWindow, Ui_MainWindow):
    def __init__(self):
        super().__init__()
        self.setupUi(self)

        # 初始化变量
        self.sections = [
            "案件卷宗信息",
            "检材信息",
            "样本信息",
            "预处理结果",
            "特征比对",
            "测量比对",
        ]
        self.template_path = None

        # 初始化顺序列表
        self.initialize_list()

        # 绑定按钮信号和槽
        self.button_move_up.clicked.connect(self.move_up)
        self.button_move_down.clicked.connect(self.move_down)
        self.button_load_template.clicked.connect(self.load_template)
        self.button_export_word.clicked.connect(self.export_word)

    def initialize_list(self):
        """初始化顺序列表"""
        self.listwidget_order.addItems(self.sections)

    def move_up(self):
        """上移选中项"""
        current_row = self.listwidget_order.currentRow()
        if current_row > 0:
            item = self.listwidget_order.takeItem(current_row)
            self.listwidget_order.insertItem(current_row - 1, item)
            self.listwidget_order.setCurrentRow(current_row - 1)

    def move_down(self):
        """下移选中项"""
        current_row = self.listwidget_order.currentRow()
        if current_row < self.listwidget_order.count() - 1:
            item = self.listwidget_order.takeItem(current_row)
            self.listwidget_order.insertItem(current_row + 1, item)
            self.listwidget_order.setCurrentRow(current_row + 1)

    def load_template(self):
        """加载模板文件"""
        file_path, _ = QFileDialog.getOpenFileName(self, "Load Template", "", "Word Files (*.docx)")
        if file_path:
            self.template_path = file_path
            QMessageBox.information(self, "Success", f"Template loaded: {file_path}")

    def export_word(self):
        """生成 Word 文档"""
        if not self.template_path:
            QMessageBox.warning(self, "Error", "Please load a template first!")
            return

        # 加载模板
        document = Document(self.template_path)

        # 按顺序插入内容
        for index in range(self.listwidget_order.count()):
            section = self.listwidget_order.item(index).text()
            document.add_heading(section, level=1)
            document.add_paragraph(f"Content for {section}.")

        # 保存文书
        file_path, _ = QFileDialog.getSaveFileName(self, "Save Word Document", "", "Word Files (*.docx)")
        if file_path:
            document.save(file_path)
            QMessageBox.information(self, "Success", f"Word document saved: {file_path}")


if __name__ == "__main__":
    app = QApplication(sys.argv)
    window = DocumentGenerationTool()
    window.show()
    sys.exit(app.exec_())
